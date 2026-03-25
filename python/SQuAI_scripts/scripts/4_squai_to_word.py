"""
4_squai_to_word.py
------------------
Converts SQuAI output (JSONL) to human-readable Word (.docx) documents,
one document per species.

Each document contains:
  - Cover page: species name, date, data sources
  - One section per FinnPRIO criterion (ENT1, EST1, IMP1, etc.)
  - Question in bold
  - Answer in plain text
  - Supporting evidence passages with source reference

Input:
  squai_corpus/
    Thrips_palmi/
      squai_results/
        *.jsonl     ← SQuAI output, one line per question

Output:
  word_reports/
    Thrips_palmi_SQuAI_report_20260324.docx
    Liriomyza_huidobrensis_SQuAI_report_20260324.docx
    ...

Dependencies:
    pip install python-docx
"""

# ==============================================================================
# USER CONFIGURATION
# ==============================================================================

# Root folder containing per-species SQuAI corpus + results.
# Must match OUTPUT_DIR in pdf_to_squai_corpus.py.
SQUAI_CORPUS_DIR = r"C:\Users\dafl\Python\SQuAI\squai_corpus"   # <-- SET THIS

# Output folder for Word documents.
WORD_OUTPUT_DIR  = r"C:\Users\dafl\Python\SQuAI\word_reports"   # <-- SET THIS

# Process only this species (sub-folder name with underscores).
# Set to None to process all species with squai_results/ present.
SPECIES_FILTER = "Thrips_palmi"   # e.g. "Thrips_palmi" or None

# Maximum number of evidence passages to include per question.
MAX_EVIDENCE = 3

# Include evidence passages in the document (True/False).
INCLUDE_EVIDENCE = True

# ==============================================================================
# END OF USER CONFIGURATION
# ==============================================================================

import json
import logging
import argparse
from pathlib import Path
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Mm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("ERROR: python-docx not installed. Run: pip install python-docx")

logging.basicConfig(level=logging.INFO, format="%(levelname)s | %(message)s")
log = logging.getLogger(__name__)

# FinnPRIO section groupings for document structure
SECTION_GROUPS = {
    "Entry":         ["ENT1", "ENT2A", "ENT2B", "ENT3", "ENT4"],
    "Establishment": ["EST1", "EST2", "EST3", "EST4"],
    "Impact":        ["IMP1", "IMP2_1", "IMP2_2", "IMP2_3", "IMP3", "IMP4"],
    "Management":    ["MAN1", "MAN2", "MAN3", "MAN4", "MAN5"],
}

# Human-readable criterion labels
CRITERION_LABELS = {
    "ENT1":   "ENT1 – Global geographical distribution",
    "ENT2A":  "ENT2A – Transport via pathway (without official measures)",
    "ENT2B":  "ENT2B – Transport via pathway (with official measures)",
    "ENT3":   "ENT3 – Volume of traded host plant commodity",
    "ENT4":   "ENT4 – Transfer to suitable habitat after entry",
    "EST1":   "EST1 – Reproduction and overwintering in Norway",
    "EST2":   "EST2 – Host plant availability in Norway",
    "EST3":   "EST3 – Spread rate within Norway",
    "EST4":   "EST4 – Characteristics assisting establishment/spread",
    "IMP1":   "IMP1 – Direct economic losses",
    "IMP2_1": "IMP2.1 – Impact on foreign trade",
    "IMP2_2": "IMP2.2 – Vector for other pests",
    "IMP2_3": "IMP2.3 – Impact on plant production sector profitability",
    "IMP3":   "IMP3 – Impact on natural ecosystems",
    "IMP4":   "IMP4 – Environmental and social impacts",
    "MAN1":   "MAN1 – Natural spread to Norway within 10 years",
    "MAN2":   "MAN2 – Presence in the European Union",
    "MAN3":   "MAN3 – Detectability during inspections",
    "MAN4":   "MAN4 – Difficulty of eradication",
    "MAN5":   "MAN5 – Difficulty of surveillance",
}


# ── helpers ────────────────────────────────────────────────────────────────────

def add_horizontal_rule(doc: "Document") -> None:
    """Add a thin horizontal line (paragraph bottom border)."""
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_paragraph_spacing(para, before: int = 0, after: int = 6) -> None:
    pPr  = para._p.get_or_add_pPr()
    spac = OxmlElement("w:spacing")
    spac.set(qn("w:before"), str(before))
    spac.set(qn("w:after"),  str(after))
    pPr.append(spac)


def load_squai_results(results_dir: Path) -> list[dict]:
    """Load all JSONL result files from squai_results/ directory."""
    records = []
    for jsonl_file in sorted(results_dir.rglob("*.jsonl")):
        with open(jsonl_file, encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if line:
                    try:
                        records.append(json.loads(line))
                    except json.JSONDecodeError as e:
                        log.warning("Skipping malformed line in %s: %s", jsonl_file.name, e)
    return records


def criterion_id_from_record(rec: dict) -> str:
    """Extract base criterion ID (e.g. 'ENT1') from record id field."""
    rid = rec.get("id", "")
    # id format: "ENT1__Thrips_palmi"  or just "ENT1"
    return rid.split("__")[0] if "__" in rid else rid


def build_results_index(records: list[dict]) -> dict[str, dict]:
    """Index records by criterion ID for fast lookup."""
    index = {}
    for rec in records:
        cid = criterion_id_from_record(rec)
        if cid:
            index[cid] = rec
    return index


# ── document builder ───────────────────────────────────────────────────────────

def build_word_document(species_name: str, results_index: dict,
                        meta: dict, output_path: Path) -> None:
    """Build and save a Word document for one species."""

    doc = Document()

    # ── page setup (A4) ────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Mm(210)   # A4
    section.page_height = Mm(297)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── styles ─────────────────────────────────────────────────────────────
    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Arial"
    style_normal.font.size = Pt(11)

    for h_name, h_size, h_bold in [
        ("Heading 1", 16, True),
        ("Heading 2", 13, True),
        ("Heading 3", 11, True),
    ]:
        s = doc.styles[h_name]
        s.font.name  = "Arial"
        s.font.size  = Pt(h_size)
        s.font.bold  = h_bold
        s.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)   # VKM-ish dark blue

    # ── cover / title ──────────────────────────────────────────────────────
    title_para = doc.add_heading(level=1)
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.runs[0] if title_para.runs else title_para.add_run()
    run.text = f"SQuAI Literature Synthesis"
    run.font.size = Pt(20)

    species_para = doc.add_paragraph()
    r = species_para.add_run(species_name.replace("_", " "))
    r.bold = True
    r.italic = True
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    doc.add_paragraph()   # spacer

    # metadata table
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_data = [
        ("Generated",      datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Species",        species_name.replace("_", " ")),
        ("PDFs indexed",   str(meta.get("n_pdfs", "–"))),
        ("Chunks indexed", str(meta.get("n_chunks", "–"))),
    ]
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── introductory note ──────────────────────────────────────────────────
    note = doc.add_paragraph()
    bold_run = note.add_run("Note: ")
    bold_run.bold = True
    bold_run.font.size = Pt(10)
    bold_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    text_run = note.add_run(
        "This document was generated automatically by the SQuAI multi-agent "
        "RAG pipeline from a curated literature corpus. Answers and evidence "
        "passages are extracted from indexed scientific literature. All content "
        "should be reviewed and verified by a qualified assessor before use in "
        "formal risk assessment."
    )
    text_run.font.size = Pt(10)
    text_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    doc.add_paragraph()

    # ── main content: one section per FinnPRIO group ───────────────────────
    for group_name, criterion_ids in SECTION_GROUPS.items():
        doc.add_heading(group_name, level=1)

        any_content = False
        for cid in criterion_ids:
            rec = results_index.get(cid)
            label = CRITERION_LABELS.get(cid, cid)

            # criterion heading
            doc.add_heading(label, level=2)

            if rec is None:
                p = doc.add_paragraph()
                p.add_run("No result available for this criterion.").italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)
                doc.add_paragraph()
                continue

            any_content = True

            # question
            q_para = doc.add_paragraph()
            q_run  = q_para.add_run("Question: ")
            q_run.bold = True
            q_run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
            q_para.add_run(rec.get("question", ""))
            set_paragraph_spacing(q_para, before=0, after=60)

            # answer
            answer = rec.get("answer", "").strip()
            if answer:
                a_para = doc.add_paragraph()
                a_run  = a_para.add_run("Answer: ")
                a_run.bold = True
                a_para.add_run(answer)
                set_paragraph_spacing(a_para, before=0, after=60)
            else:
                p = doc.add_paragraph()
                p.add_run("No answer generated.").italic = True
                p.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

            # evidence passages
            if INCLUDE_EVIDENCE:
                evidence = rec.get("evidence", [])[:MAX_EVIDENCE]
                if evidence:
                    ev_heading = doc.add_paragraph()
                    ev_heading.add_run("Supporting evidence:").bold = True
                    ev_heading.runs[0].font.size = Pt(10)

                    for ev in evidence:
                        sentence = ev.get("sentence", ev.get("text", "")).strip()
                        paper_id = ev.get("paper_id", "")
                        score    = ev.get("score", "")
                        if not sentence:
                            continue
                        ev_para = doc.add_paragraph(style="List Bullet")
                        ev_para.add_run(f'"{sentence}"')
                        ev_para.runs[0].font.size = Pt(10)
                        ev_para.runs[0].italic = True
                        if paper_id or score:
                            src_text = f"  [id: {paper_id}"
                            if score:
                                src_text += f", score: {float(score):.2f}"
                            src_text += "]"
                            src_run = ev_para.add_run(src_text)
                            src_run.font.size = Pt(9)
                            src_run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

            doc.add_paragraph()  # spacer between criteria

        if not any_content:
            p = doc.add_paragraph()
            p.add_run("No SQuAI results found for this section.").italic = True

        add_horizontal_rule(doc)
        doc.add_paragraph()

    # ── footer note ────────────────────────────────────────────────────────
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run  = footer_para.add_run(
        f"Generated by SQuAI BioPRIO pipeline  ·  "
        f"{datetime.now().strftime('%Y-%m-%d')}  ·  VKM"
    )
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # ── save ───────────────────────────────────────────────────────────────
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    log.info("[%s] Saved → %s", species_name, output_path)


# ── main ───────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Convert SQuAI results to human-readable Word documents"
    )
    parser.add_argument(
        "--corpus_dir", default=SQUAI_CORPUS_DIR,
        help=f"SQuAI corpus root directory (default: {SQUAI_CORPUS_DIR})"
    )
    parser.add_argument(
        "--output_dir", default=WORD_OUTPUT_DIR,
        help=f"Output directory for Word documents (default: {WORD_OUTPUT_DIR})"
    )
    parser.add_argument(
        "--species", default=SPECIES_FILTER,
        help="Process only this species (default: all)"
    )
    args = parser.parse_args()

    if not HAS_DOCX:
        log.error("python-docx not installed. Run: pip install python-docx")
        return

    corpus_dir = Path(args.corpus_dir)
    output_dir = Path(args.output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    # collect species to process
    # Build a mapping: species_dir -> results_dir
    # Strategy 1: find squai_results/ directories anywhere under corpus_dir (recursive)
    # Strategy 2: for species dirs that have result-like JSONL files directly inside them
    species_results: dict[Path, Path] = {}

    if args.species:
        sp_dir = corpus_dir / args.species
        # search recursively for squai_results under this species dir
        for rd in sorted(sp_dir.rglob("squai_results")):
            if rd.is_dir():
                species_results[sp_dir] = rd
                break
        else:
            # fallback: use the species dir itself
            species_results[sp_dir] = sp_dir
    else:
        # Find all squai_results/ dirs anywhere under corpus_dir
        for rd in sorted(corpus_dir.rglob("squai_results")):
            if rd.is_dir():
                # The species dir is the first-level subdir of corpus_dir
                # that contains this squai_results dir
                parts = rd.relative_to(corpus_dir).parts
                sp_dir = corpus_dir / parts[0]
                species_results.setdefault(sp_dir, rd)

        # If nothing found, fall back to any first-level subdir that has JSONL files
        # (excluding corpus.jsonl) — results may have been saved directly in the species dir
        if not species_results:
            for sp_dir in sorted(d for d in corpus_dir.iterdir() if d.is_dir()):
                jsonl_files = [
                    f for f in sp_dir.rglob("*.jsonl")
                    if f.name != "corpus.jsonl"
                ]
                if jsonl_files:
                    species_results[sp_dir] = sp_dir

    if not species_results:
        log.error(
            "No results found (searched recursively) in %s. "
            "Run pdf_to_squai_corpus.py --run_squai first.", corpus_dir
        )
        return

    log.info("Generating Word documents for %d species…", len(species_results))
    date_str = datetime.now().strftime("%Y%m%d")

    for sp_dir, results_dir in sorted(species_results.items()):
        species_name = sp_dir.name

        # load results
        records = load_squai_results(results_dir)
        if not records:
            log.warning("[%s] No results found in %s – skipping", species_name, results_dir)
            continue

        results_index = build_results_index(records)
        log.info("[%s] %d criteria with results", species_name, len(results_index))

        # load meta
        meta_path = sp_dir / "meta.json"
        meta = {}
        if meta_path.exists():
            with open(meta_path, encoding="utf-8") as f:
                meta = json.load(f)

        # output path
        out_path = output_dir / f"{species_name}_SQuAI_report_{date_str}.docx"
        build_word_document(species_name, results_index, meta, out_path)

    log.info("Done. Word documents written to %s", output_dir)


if __name__ == "__main__":
    main()
